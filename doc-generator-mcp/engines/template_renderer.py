"""模板渲染引擎 —— Word 文档生成核心"""

import shutil
from pathlib import Path
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def render_word_from_template(
    template_path: str,
    output_path: str,
    context: dict,
) -> str:
    """使用 docxtpl 渲染 Word 模板"""
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    tpl.save(output_path)
    return output_path


def render_word_from_scratch(
    output_path: str,
    context: dict,
) -> str:
    """无模板时，使用 python-docx 从零创建 Word 文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(11)

    # 标题
    title = context.get("title", "未命名文档")
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 日期
    date_str = context.get("date", "")
    if date_str:
        date_para = doc.add_paragraph(date_str)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 章节
    for section in context.get("sections", []):
        doc.add_heading(section["heading"], level=1)
        for para_text in section.get("content", []):
            doc.add_paragraph(para_text)

    # 表格（兼容 "tables" 和 "table_data" 两种 key）
    tables = context.get("tables") or context.get("table_data") or []
    for table_data in tables:
        if table_data.get("caption"):
            doc.add_paragraph(table_data["caption"], style="Caption")

        headers = table_data["headers"]
        rows = table_data.get("rows", [])

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"

        # 表头
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = str(h)
            for run in cell.paragraphs[0].runs:
                run.bold = True

        # 数据行
        for row_idx, row_data in enumerate(rows):
            for col_idx, val in enumerate(row_data):
                table.cell(row_idx + 1, col_idx).text = str(val)

    doc.save(output_path)
    return output_path


def get_template_context(schema_dict: dict) -> dict:
    """将 Schema 转换为模板渲染所需的 context"""
    return {
        "title": schema_dict.get("title", ""),
        "date": schema_dict.get("date", ""),
        "sections": [
            {
                "heading": s["heading"],
                "content": s.get("content", []),
            }
            for s in schema_dict.get("sections", [])
        ],
        "table_data": [
            {
                "caption": t.get("caption", ""),
                "headers": t["headers"],
                "rows": t.get("rows", []),
            }
            for t in schema_dict.get("tables", [])
        ],
    }

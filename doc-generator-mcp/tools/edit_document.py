"""文档编辑工具 —— 定向修改已有文档"""

from pathlib import Path
from utils.backup import smart_backup
from schemas.edit_schema import EditInstruction, EditResult


def edit_document(
    file_path: str,
    edits: list[EditInstruction],
    backup: bool = True,
    output_path: str = None,
) -> dict:
    """对已有文档执行定向编辑

    Args:
        file_path: 文档文件路径
        edits: 编辑指令列表
        backup: 是否自动备份
        output_path: 输出路径（不填则覆盖原文件）

    Returns:
        编辑结果
    """
    path = Path(file_path)
    if not path.exists():
        return EditResult(success=False, errors=[f"文件不存在: {file_path}"]).model_dump()

    ext = path.suffix.lower()
    if ext == ".docx":
        return _edit_word(file_path, edits, backup, output_path)
    elif ext == ".xlsx":
        return _edit_excel(file_path, edits, backup, output_path)
    else:
        return EditResult(success=False, errors=[f"不支持的文件类型: {ext}"]).model_dump()


def _edit_word(file_path: str, edits: list[EditInstruction], backup: bool, output_path: str) -> dict:
    """编辑 Word 文档"""
    from docx import Document

    path = Path(file_path)
    target = Path(output_path) if output_path else path

    # 智能备份
    backup_path = smart_backup(file_path, backup, output_path)

    doc = Document(str(path))
    changes = []
    errors = []

    for edit in edits:
        try:
            if edit.action == "replace_text":
                count = _replace_text_in_doc(doc, edit.search, edit.replace, edit.scope)
                changes.append(f"替换文本 '{edit.search}' → '{edit.replace}' ({count} 处)")

            elif edit.action == "update_table":
                if edit.table_index is not None and edit.data is not None:
                    table = doc.tables[edit.table_index]
                    for row_idx, row_data in enumerate(edit.data):
                        for col_idx, val in enumerate(row_data):
                            target_row = edit.start_row + row_idx
                            if target_row < len(table.rows) and col_idx < len(table.columns):
                                cell = table.cell(target_row, col_idx)
                                # 保留格式，只改文本
                                if cell.paragraphs[0].runs:
                                    cell.paragraphs[0].runs[0].text = str(val)
                                else:
                                    cell.text = str(val)
                    changes.append(f"更新表格 {edit.table_index} ({len(edit.data)} 行)")

            elif edit.action == "insert_paragraph":
                if edit.after_index is not None and edit.text is not None:
                    ref_para = doc.paragraphs[edit.after_index]
                    # 在参考段落之后插入新段落，自动复制完整格式
                    _insert_paragraph_after(ref_para, edit.text)
                    changes.append(f"在第 {edit.after_index} 段后插入新段落")

            elif edit.action == "update_paragraph_format":
                if edit.index is not None:
                    para = doc.paragraphs[edit.index]
                    _update_paragraph_format(
                        para,
                        alignment=edit.alignment,
                        space_before=edit.space_before,
                        space_after=edit.space_after,
                        line_spacing=edit.line_spacing,
                        first_line_indent=edit.first_line_indent,
                        left_indent=edit.left_indent,
                        right_indent=edit.right_indent,
                    )
                    changes.append(f"更新第 {edit.index} 段格式")

            elif edit.action == "delete_paragraph":
                if edit.index is not None:
                    para = doc.paragraphs[edit.index]
                    para._element.getparent().remove(para._element)
                    changes.append(f"删除第 {edit.index} 段")

            elif edit.action == "update_cell":
                if edit.table_index is not None and edit.row is not None and edit.column is not None:
                    table = doc.tables[edit.table_index]
                    cell = table.cell(edit.row, edit.column)
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].text = str(edit.value)
                    else:
                        cell.text = str(edit.value)
                    changes.append(f"更新表格 {edit.table_index} 单元格 ({edit.row},{edit.column})")

        except Exception as e:
            errors.append(f"编辑指令 {edit.action.value} 执行失败: {e}")

    # 保存
    doc.save(str(target))

    return EditResult(
        success=len(errors) == 0,
        file_path=str(target),
        backup_path=backup_path,
        changes_applied=len(changes),
        summary="; ".join(changes) if changes else "无变更",
        errors=errors,
    ).model_dump()


def _edit_excel(file_path: str, edits: list[EditInstruction], backup: bool, output_path: str) -> dict:
    """编辑 Excel 文件"""
    from openpyxl import load_workbook

    path = Path(file_path)
    target = Path(output_path) if output_path else path

    # 智能备份
    backup_path = smart_backup(file_path, backup, output_path)

    wb = load_workbook(str(path))
    ws = wb.active
    changes = []
    errors = []

    for edit in edits:
        try:
            if edit.action == "update_cell":
                if edit.row is not None and edit.column is not None:
                    cell = ws.cell(row=edit.row + 1, column=edit.column + 1)  # openpyxl 1-based
                    cell.value = edit.value
                    changes.append(f"更新单元格 ({edit.row},{edit.column})")

            elif edit.action == "update_table":
                if edit.data is not None:
                    for row_idx, row_data in enumerate(edit.data):
                        for col_idx, val in enumerate(row_data):
                            r = edit.start_row + row_idx + 1  # 1-based
                            cell = ws.cell(row=r, column=col_idx + 1)
                            cell.value = val
                    changes.append(f"更新数据区域 ({len(edit.data)} 行)")

            elif edit.action == "insert_row":
                if edit.at_row is not None:
                    ws.insert_rows(edit.at_row + 1)
                    changes.append(f"在第 {edit.at_row} 行后插入新行")

        except Exception as e:
            errors.append(f"编辑指令 {edit.action.value} 执行失败: {e}")

    wb.save(str(target))
    wb.close()

    return EditResult(
        success=len(errors) == 0,
        file_path=str(target),
        backup_path=backup_path,
        changes_applied=len(changes),
        summary="; ".join(changes) if changes else "无变更",
        errors=errors,
    ).model_dump()


# ========== 辅助函数 ==========

def _replace_text_in_doc(doc, search: str, replace: str, scope: str) -> int:
    """在文档中替换文本，保留格式"""
    count = 0
    for para in doc.paragraphs:
        if search in para.text:
            for run in para.runs:
                if search in run.text:
                    run.text = run.text.replace(search, replace)
                    count += 1
    # 表格内搜索
    if scope == "all":
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if search in para.text:
                            for run in para.runs:
                                if search in run.text:
                                    run.text = run.text.replace(search, replace)
                                    count += 1
    return count


def _insert_paragraph_after(reference, text):
    """在参考段落之后插入新段落（纯 XML 操作，避免 parent.part 问题）

    自动复制参考段落的完整 pPr（段落属性），包括样式、字体、字号、
    粗体、斜体、颜色、对齐等所有格式信息。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy

    new_p = OxmlElement("w:p")

    # 复制参考段落的完整 pPr（段落属性）
    ref_pPr = reference._element.find(qn("w:pPr"))
    if ref_pPr is not None:
        new_p.append(deepcopy(ref_pPr))

    # 创建文本 run，并复制参考段落首个 run 的字符格式（rPr）
    new_r = OxmlElement("w:r")

    # 复制参考段落首个 run 的字符格式
    ref_runs = reference._element.findall(qn("w:r"))
    if ref_runs:
        ref_rPr = ref_runs[0].find(qn("w:rPr"))
        if ref_rPr is not None:
            new_r.append(deepcopy(ref_rPr))

    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set(qn("xml:space"), "preserve")
    new_r.append(new_t)
    new_p.append(new_r)
    reference._element.addnext(new_p)
    return new_p


def _update_paragraph_format(
    paragraph,
    alignment=None,
    space_before=None,
    space_after=None,
    line_spacing=None,
    first_line_indent=None,
    left_indent=None,
    right_indent=None,
):
    """更新段落格式（对齐、间距、缩进），保留其他已有格式"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._element.insert(0, pPr)

    # 对齐方式
    if alignment is not None:
        alignment_map = {
            "left": "left",
            "center": "center",
            "right": "right",
            "justify": "both",
        }
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), alignment_map.get(alignment, alignment))

    # 间距和缩进通过 w:spacing / w:ind 元素设置
    if any(v is not None for v in [space_before, space_after, line_spacing]):
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        if space_before is not None:
            spacing.set(qn("w:before"), str(int(space_before * 20)))  # pt -> twips
        if space_after is not None:
            spacing.set(qn("w:after"), str(int(space_after * 20)))
        if line_spacing is not None:
            spacing.set(qn("w:line"), str(int(line_spacing * 240)))
            spacing.set(qn("w:lineRule"), "auto")

    if any(v is not None for v in [first_line_indent, left_indent, right_indent]):
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        if left_indent is not None:
            ind.set(qn("w:left"), str(int(left_indent * 567)))  # cm -> twips
        if right_indent is not None:
            ind.set(qn("w:right"), str(int(right_indent * 567)))
        if first_line_indent is not None:
            ind.set(qn("w:firstLine"), str(int(first_line_indent * 567)))

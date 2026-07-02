"""文档读取工具 —— 读取已有文档的结构和内容"""

from pathlib import Path


def read_document(
    file_path: str,
    include_content: bool = True,
    include_style: bool = False,
    max_depth: int = 2,
) -> dict:
    """读取已有文档的结构和内容

    Args:
        file_path: 文档文件路径
        include_content: 是否包含段落文本内容
        include_style: 是否包含样式详情
        max_depth: 结构解析深度 (1=仅标题, 2=标题+段落, 3=完整)

    Returns:
        文档结构描述 JSON
    """
    path = Path(file_path)
    if not path.exists():
        return {"success": False, "error": f"文件不存在: {file_path}"}

    ext = path.suffix.lower()

    if ext == ".docx":
        return _read_word(file_path, include_content, include_style, max_depth)
    elif ext == ".xlsx":
        return _read_excel(file_path, include_content, max_depth)
    else:
        return {"success": False, "error": f"不支持的文件类型: {ext}"}


def _read_word(file_path: str, include_content: bool, include_style: bool, max_depth: int) -> dict:
    """读取 Word 文档结构"""
    from docx import Document

    doc = Document(file_path)

    sections = []
    current_section = None
    para_index = 0

    for para in doc.paragraphs:
        style_name = para.style.name

        # 检测标题（新章节开始）
        if style_name.startswith("Heading"):
            if current_section is not None:
                sections.append(current_section)

            current_section = {
                "index": para_index,
                "heading": para.text,
                "style": style_name,
                "paragraphs": [],
            }
            if include_style:
                current_section["style_detail"] = {
                    "font_name": para.style.font.name if para.style.font else None,
                    "font_size": str(para.style.font.size) if para.style.font and para.style.font.size else None,
                }
        elif current_section is not None and max_depth >= 2:
            # 段落属于当前章节
            para_info = {"index": para_index, "text": para.text if include_content else ""}
            if include_style:
                para_info["style"] = style_name
            current_section["paragraphs"].append(para_info)
        elif current_section is None and max_depth >= 2:
            # 标题前的内容（如文档标题、日期）
            pass

        para_index += 1

    # 添加最后一个章节
    if current_section is not None:
        sections.append(current_section)

    # 表格信息
    tables = []
    for i, table in enumerate(doc.tables):
        table_info = {
            "index": i,
            "rows": len(table.rows),
            "cols": len(table.columns),
        }
        if include_content and max_depth >= 3:
            table_info["data"] = [
                [cell.text for cell in row.cells]
                for row in table.rows
            ]
        tables.append(table_info)

    return {
        "success": True,
        "file_type": "word",
        "file_path": file_path,
        "structure": {
            "sections": sections,
            "tables": tables,
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
        },
    }


def _read_excel(file_path: str, include_content: bool, max_depth: int) -> dict:
    """读取 Excel 文件结构"""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, data_only=True)

    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_info = {
            "name": sheet_name,
            "max_row": ws.max_row,
            "max_column": ws.max_column,
            "merged_cells": [str(m) for m in ws.merged_cells.ranges],
        }

        if include_content and max_depth >= 2:
            # 读取前 20 行数据（避免大文件）
            rows_data = []
            for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 20), values_only=True):
                rows_data.append([str(c) if c is not None else "" for c in row])
            sheet_info["preview_data"] = rows_data

        sheets.append(sheet_info)

    wb.close()

    return {
        "success": True,
        "file_type": "excel",
        "file_path": file_path,
        "structure": {
            "sheets": sheets,
            "total_sheets": len(sheets),
        },
    }

"""Word 嵌入图片替换工具 —— 替换 Word 文档中嵌入的图片/OLE 对象"""

import os
import shutil
import logging
from pathlib import Path
from utils.backup import smart_backup

logger = logging.getLogger(__name__)


def replace_image(
    file_path: str,
    new_image_path: str,
    target: str = None,
    backup: bool = True,
    output_path: str = None,
) -> dict:
    """替换 Word 文档中嵌入的图片

    定位方式（按优先级）：
    1. target="图 1" → 按图题文本定位附近图片
    2. target="rId14" → 按 relationship ID 直接定位
    3. target=0 → 按图片索引（第 N 个图片）

    Args:
        file_path: Word 文档路径
        new_image_path: 新图片文件路径
        target: 定位方式（图题文本/rId/索引）
        backup: 是否备份
        output_path: 输出路径

    Returns:
        替换结果
    """
    from docx import Document
    from docx.oxml.ns import qn

    if not Path(file_path).exists():
        return {"success": False, "error": f"文件不存在: {file_path}"}
    if not Path(new_image_path).exists():
        return {"success": False, "error": f"图片文件不存在: {new_image_path}"}

    # 备份
    backup_path = smart_backup(file_path, backup, output_path)

    target_path = output_path or file_path
    if output_path:
        shutil.copy2(file_path, output_path)

    try:
        doc = Document(target_path)

        # 读取新图片数据
        with open(new_image_path, "rb") as f:
            new_data = f.read()

        # 确定新图片的 content type
        ext = Path(new_image_path).suffix.lower()
        content_type_map = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
            ".emf": "image/x-emf",
            ".wmf": "image/x-wmf",
            ".svg": "image/svg+xml",
            ".tiff": "image/tiff",
        }
        new_content_type = content_type_map.get(ext, "image/png")

        # 定位目标图片的 rId
        target_rids = _locate_image(doc, target)

        if not target_rids:
            return {
                "success": False,
                "error": f"未找到目标图片 (target='{target}')",
                "backup_path": backup_path,
                "hint": "可用 target 方式: 图题文本(如'图 1')、rId(如'rId14')、图片索引(如 0)",
            }

        # 替换图片（过滤掉非图片类型）
        replaced = []
        for rid in target_rids:
            try:
                part = doc.part.related_parts[rid]
                # 检查是否为图片/OLE类型（排除页脚/页眉等非图片引用）
                ct = part.content_type.lower()
                # OLE 对象警告
                is_ole = any(t in ct for t in ["visio", "ms-visio", "oleobject"])
                if is_ole:
                    replaced.append({
                        "rId": rid,
                        "warning": "目标是 Visio OLE 对象，替换为图片会丢失可编辑性。建议使用 embed_visio_to_word 工具。",
                        "skipped": True,
                    })
                    continue
                is_image_type = any(t in ct for t in [
                    "image/", "visio", "oleobject", "ms-visio", "emf", "wmf",
                ])
                if not is_image_type:
                    replaced.append({"rId": rid, "skipped": True, "reason": f"非图片类型: {ct}"})
                    continue

                old_type = part.content_type
                old_size = len(part.blob)
                part._blob = new_data
                part._content_type = new_content_type
                replaced.append({
                    "rId": rid,
                    "old_type": old_type,
                    "new_type": new_content_type,
                    "old_size": old_size,
                    "new_size": len(new_data),
                })
            except Exception as e:
                replaced.append({"rId": rid, "error": str(e)})

        # 保存
        doc.save(target_path)

        return {
            "success": True,
            "file_path": target_path,
            "backup_path": backup_path,
            "replaced": replaced,
            "replaced_count": len([r for r in replaced if "error" not in r]),
        }

    except Exception as e:
        return {"success": False, "error": f"替换图片失败: {e}", "backup_path": backup_path}


def list_images(file_path: str) -> dict:
    """列出 Word 文档中所有嵌入的图片

    Args:
        file_path: Word 文档路径

    Returns:
        图片列表
    """
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree
    import re

    if not Path(file_path).exists():
        return {"success": False, "error": f"文件不存在: {file_path}"}

    try:
        doc = Document(file_path)
        images = []

        # 遍历所有段落，查找图片
        embed_pattern = re.compile(r'(?:r:embed|r:id)="(rId\d+)"')
        for para_idx, para in enumerate(doc.paragraphs):
            xml_str = etree.tostring(para._element, encoding="unicode")
            refs = embed_pattern.findall(xml_str)

            if refs:
                # 查找附近的图题
                caption = ""
                for j in range(para_idx, min(para_idx + 3, len(doc.paragraphs))):
                    text = doc.paragraphs[j].text.strip()
                    if text.startswith("图"):
                        caption = text
                        break

                for rid in refs:
                    img_info = {"paragraph_index": para_idx, "rId": rid, "caption": caption}

                    # 获取图片详情
                    if rid in doc.part.rels:
                        rel = doc.part.rels[rid]
                        img_info["target_ref"] = rel.target_ref
                        img_info["content_type"] = rel.target_part.content_type
                        img_info["size_bytes"] = len(rel.target_part.blob)

                    images.append(img_info)

        # 也列出所有图片 relationships
        all_image_rels = []
        for rid, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                all_image_rels.append({
                    "rId": rid,
                    "content_type": rel.target_part.content_type,
                    "size_bytes": len(rel.target_part.blob),
                    "target_ref": rel.target_ref,
                })

        return {
            "success": True,
            "file_path": file_path,
            "inline_images": images,
            "all_image_relationships": all_image_rels,
            "total_inline": len(images),
            "total_relationships": len(all_image_rels),
        }

    except Exception as e:
        return {"success": False, "error": f"列出图片失败: {e}"}


def _locate_image(doc, target) -> list:
    """根据 target 定位图片的 rId 列表"""
    from docx.oxml.ns import qn
    from lxml import etree
    import re

    if target is None:
        return []

    # 方式 1: rId 直接定位
    if isinstance(target, str) and target.startswith("rId"):
        if target in doc.part.rels:
            return [target]
        return []

    # 方式 2: 图片索引
    if isinstance(target, int):
        all_rids = []
        for para in doc.paragraphs:
            xml_str = etree.tostring(para._element, encoding="unicode")
            # 匹配 r:embed (DrawingML) 和 r:id (VML/OLE)
            refs = re.findall(r'(?:r:embed|r:id)="(rId\d+)"', xml_str)
            all_rids.extend(refs)
        if 0 <= target < len(all_rids):
            return [all_rids[target]]
        return []

    # 方式 3: 图题文本定位
    if isinstance(target, str):
        # 找到包含 target 文本的段落
        target_para_idx = None
        for i, para in enumerate(doc.paragraphs):
            if target in para.text:
                target_para_idx = i
                break

        if target_para_idx is None:
            return []

        # 在图题前后搜索图片（向前搜索范围更大）
        found_rids = []
        embed_pattern = re.compile(r'(?:r:embed|r:id)="(rId\d+)"')
        for i in range(max(0, target_para_idx - 15), min(len(doc.paragraphs), target_para_idx + 3)):
            para = doc.paragraphs[i]
            xml_str = etree.tostring(para._element, encoding="unicode")
            refs = embed_pattern.findall(xml_str)
            if refs and i <= target_para_idx:
                found_rids.extend(refs)

        return found_rids

    return []

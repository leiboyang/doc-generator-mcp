"""分级验证器 —— 文档质量保障核心"""

from pathlib import Path
from dataclasses import dataclass, field


@dataclass
class ValidationResult:
    """分级验证结果"""
    critical_errors: list[str] = field(default_factory=list)
    important_errors: list[str] = field(default_factory=list)
    suggestions: list[str] = field(default_factory=list)

    @property
    def passed(self) -> bool:
        """只有 critical 错误才阻止输出"""
        return len(self.critical_errors) == 0

    @property
    def quality_score(self) -> int:
        """质量评分 0-100"""
        total = (
            len(self.critical_errors) * 20
            + len(self.important_errors) * 5
            + len(self.suggestions) * 1
        )
        return max(0, 100 - total)

    def to_dict(self) -> dict:
        return {
            "passed": self.passed,
            "quality_score": self.quality_score,
            "critical_errors": self.critical_errors,
            "important_errors": self.important_errors,
            "suggestions": self.suggestions,
        }


def validate_word_document(file_path: str, expected_schema: dict = None) -> ValidationResult:
    """验证 Word 文档"""
    from docx import Document

    result = ValidationResult()
    path = Path(file_path)

    # CRITICAL: 文件存在且非空
    if not path.exists():
        result.critical_errors.append(f"文件不存在: {file_path}")
        return result

    if path.stat().st_size == 0:
        result.critical_errors.append("文件为空")
        return result

    try:
        doc = Document(str(path))
    except Exception as e:
        result.critical_errors.append(f"无法打开文档: {e}")
        return result

    # CRITICAL: 文档不为空
    if len(doc.paragraphs) == 0:
        result.critical_errors.append("文档没有任何段落")
        return result

    if expected_schema:
        # CRITICAL: 标题存在
        title = expected_schema.get("title", "")
        if title:
            title_found = any(title in p.text for p in doc.paragraphs[:5])
            if not title_found:
                result.critical_errors.append(f"标题 '{title}' 未在文档前部找到")

        # IMPORTANT: 章节完整性
        expected_sections = expected_schema.get("sections", [])
        found_headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        for section in expected_sections:
            heading = section.get("heading", "")
            if heading and not any(heading in h for h in found_headings):
                result.important_errors.append(f"缺失章节: {heading}")

        # IMPORTANT: 段落数在合理范围
        expected_para_count = sum(
            len(s.get("content", [])) + 1 for s in expected_sections
        )
        actual_count = len(doc.paragraphs)
        if expected_para_count > 0 and actual_count < expected_para_count * 0.6:
            result.important_errors.append(
                f"段落数偏少: 预期 ~{expected_para_count}, 实际 {actual_count}"
            )

        # IMPORTANT: 表格数量
        expected_table_count = len(expected_schema.get("tables", []))
        actual_table_count = len(doc.tables)
        if actual_table_count < expected_table_count:
            result.important_errors.append(
                f"表格数不足: 预期 {expected_table_count}, 实际 {actual_table_count}"
            )

    # SUGGESTION: 样式一致性
    body_styles = set()
    for p in doc.paragraphs:
        if p.style.name not in ("Title", "Heading 1", "Heading 2", "Heading 3"):
            body_styles.add(p.style.name)
    if len(body_styles) > 4:
        result.suggestions.append(
            f"正文使用了 {len(body_styles)} 种不同样式 ({body_styles})，建议统一"
        )

    return result


def validate_excel_document(file_path: str, expected_schema: dict = None) -> ValidationResult:
    """验证 Excel 文档"""
    from openpyxl import load_workbook

    result = ValidationResult()
    path = Path(file_path)

    if not path.exists():
        result.critical_errors.append(f"文件不存在: {file_path}")
        return result

    if path.stat().st_size == 0:
        result.critical_errors.append("文件为空")
        return result

    try:
        wb = load_workbook(str(path), data_only=True)
    except Exception as e:
        result.critical_errors.append(f"无法打开文件: {e}")
        return result

    ws = wb.active

    # CRITICAL: 有数据
    if ws.max_row <= 1 and ws.max_column <= 1:
        result.critical_errors.append("工作表为空")
        return result

    if expected_schema:
        data = expected_schema.get("data", {})
        expected_rows = len(data.get("rows", []))
        actual_rows = ws.max_row - 1  # 减去表头

        # IMPORTANT: 行数
        if actual_rows < expected_rows:
            result.important_errors.append(
                f"数据行数不足: 预期 {expected_rows}, 实际 {actual_rows}"
            )

        # IMPORTANT: 抽样验证前3行数据
        for i, expected_row in enumerate(data.get("rows", [])[:3]):
            for j, expected_val in enumerate(expected_row):
                cell = ws.cell(row=i + 2, column=j + 1)
                if str(cell.value) != str(expected_val):
                    result.important_errors.append(
                        f"单元格 ({i+2},{j+1}) 值不匹配: 预期 '{expected_val}', 实际 '{cell.value}'"
                    )
                    break

    wb.close()
    return result


def validate_document(file_path: str, expected_schema: dict = None) -> dict:
    """统一验证入口"""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".docx":
        result = validate_word_document(file_path, expected_schema)
    elif ext == ".xlsx":
        result = validate_excel_document(file_path, expected_schema)
    else:
        result = ValidationResult()
        result.critical_errors.append(f"不支持的文件类型: {ext}")

    output = result.to_dict()
    output["file_path"] = file_path
    output["file_type"] = ext
    return output
